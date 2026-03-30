import json
import pandas as pd
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

# ----------------------------
# Load JSON
# ----------------------------
file_path = Path("data/model_outputs/assumption_report.json")

with open(file_path, "r") as f:
    report = json.load(f)

# ----------------------------
# Formatting helpers
# ----------------------------
def pct(x):
    return f"{100 * x:.2f}%"

def pval_fmt(x):
    if x is None:
        return ""
    if x < 0.001:
        return f"{x:.2e}"
    return f"{x:.4f}"

def num_fmt(x):
    if x is None:
        return ""
    return f"{x:.4f}"

def yes_no(flag):
    return "Yes" if flag else "No"

# ----------------------------
# Build DataFrames
# ----------------------------
missingness_df = pd.DataFrame(
    list(report["missingness"].items()),
    columns=["Variable", "Missing Proportion"]
)
missingness_df["Missing Percent"] = missingness_df["Missing Proportion"].apply(pct)
missingness_df = missingness_df.drop(columns=["Missing Proportion"])

class_map = {"1": "Warning", "2": "Citation", "3": "Arrest"}
class_balance_df = pd.DataFrame(
    list(report["class_balance"].items()),
    columns=["Outcome Code", "Proportion"]
)
class_balance_df["Outcome"] = class_balance_df["Outcome Code"].map(class_map)
class_balance_df["Percent"] = class_balance_df["Proportion"].apply(pct)
class_balance_df = class_balance_df[["Outcome Code", "Outcome", "Percent"]]

vif_df = pd.DataFrame(
    list(report["vif_table"].items()),
    columns=["Term", "VIF"]
)
vif_df["VIF"] = vif_df["VIF"].apply(num_fmt)

linearity_rows = []
for comp, vals in report["linearity_details"].items():
    linearity_rows.append({
        "Comparison": comp,
        "Wealth log p-value": pval_fmt(vals.get("wealth_log_p")),
        "Age log p-value": pval_fmt(vals.get("age_log_p"))
    })
linearity_df = pd.DataFrame(linearity_rows)

summary_df = pd.DataFrame([
    ["Zero-count race × sex cells", report["zero_count_race_sex_cells"], "No issue" if report["zero_count_race_sex_cells"] == 0 else "Sparse cells"],
    ["Maximum VIF", num_fmt(report["max_vif"]), "Acceptable multicollinearity"],
    ["Age nonlinearity", yes_no(report["age_nonlinearity_flag"]), "Spline recommended" if report["age_nonlinearity_flag"] else "Linear OK"],
    ["Wealth nonlinearity", yes_no(report["wealth_nonlinearity_flag"]), "Spline recommended" if report["wealth_nonlinearity_flag"] else "Linear OK"],
    ["IIA relative change (wealth)", num_fmt(report["iia_relative_change"]["wealth_c"]), "Check sensitivity"],
    ["IIA flag", yes_no(report["iia_flag"]), "Caution" if report["iia_flag"] else "No major concern"]
], columns=["Assumption", "Result", "Interpretation"])

recommendations_df = pd.DataFrame({
    "Recommendation": report["recommendations"]
})

# ----------------------------
# Create Word Document
# ----------------------------
doc = Document()

# Title
title = doc.add_paragraph()
run = title.add_run("Model Assumption Diagnostics Report")
run.bold = True
run.font.size = Pt(20)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("")

# ----------------------------
# Helper to add table
# ----------------------------
def add_table_from_df(document, df, title):
    # Section heading
    heading = document.add_paragraph()
    run = heading.add_run(title)
    run.bold = True
    run.font.size = Pt(14)

    document.add_paragraph("")

    table = document.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)

    # Data rows
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)

    document.add_paragraph("")

# ----------------------------
# Add sections
# ----------------------------
add_table_from_df(doc, missingness_df, "1. Missingness Summary")
add_table_from_df(doc, class_balance_df, "2. Outcome Class Balance")
add_table_from_df(doc, vif_df, "3. Variance Inflation Factors")
add_table_from_df(doc, linearity_df, "4. Linearity Diagnostics")
add_table_from_df(doc, summary_df, "5. Assumption Summary")
add_table_from_df(doc, recommendations_df, "6. Recommendations")

# Independence note
doc.add_paragraph("7. Independence Considerations", style=None).runs[0].bold = True
doc.add_paragraph(report["independence_note"])

# ----------------------------
# Save document
# ----------------------------
output_path = Path("Assumption_Report.docx")
doc.save(output_path)

print(f"Report saved to: {output_path.resolve()}")